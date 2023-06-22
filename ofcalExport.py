import openpyxl
from docx import Document

# Load the Excel file
workbook = openpyxl.load_workbook("input.xlsx")
worksheet = workbook.active

# Extract the column values from the Excel data
rows = list(worksheet.iter_rows(values_only=True))
column_names = [cell for cell in rows[0]]
data_rows = rows[1:]

# Load the Word document template
template = Document("template.docx")

# Iterate over each data row
for data_row in data_rows:
    # Create a new document based on the template
    doc = Document("template.docx")

    # Replace the placeholders in the document with column values
    for column_name, column_value in zip(column_names, data_row):
        placeholder = "{{" + column_name + "}}"
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(column_value))

    # Save the resulting document
    output_file_name = "_".join(str(column_value)[:1] for column_value in data_row)

    doc.save(f"./exports/output_{output_file_name}.docx")

print("Word documents generated successfully!")
